from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "results" / "tissuePMHC_model_principles_guide.docx"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
PAGE_WIDTH_DXA = 9360


def set_run_font(run, size=None, bold=None, color=None, italic=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col in list(grid):
        grid.remove(grid_col)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in (("top", "80"), ("bottom", "80"), ("start", "120"), ("end", "120")):
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), value)
                node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def paragraph(doc, text="", *, size=11, bold=False, color=INK, before=0, after=6, line=1.25, align=None, italic=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color, italic=italic)
    return p


def labeled_paragraph(doc, label, text):
    p = paragraph(doc, "", after=5)
    r = p.add_run(label)
    set_run_font(r, size=11, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_run_font(r, size=11, color=INK)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if level == 1:
        pf.space_before, pf.space_after = Pt(18), Pt(10)
        size, color = 16, BLUE
    elif level == 2:
        pf.space_before, pf.space_after = Pt(14), Pt(7)
        size, color = 13, BLUE
    else:
        pf.space_before, pf.space_after = Pt(10), Pt(5)
        size, color = 12, DARK_BLUE
    r = p.add_run(text)
    set_run_font(r, size=size, bold=True, color=color)
    return p


def callout(doc, title, body, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [PAGE_WIDTH_DXA])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=DARK_BLUE)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(body)
    set_run_font(r, size=10.5, color=INK)
    paragraph(doc, "", after=3)


def fill_table_cell(cell, text, *, header=False, center=False, size=9.4):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run_font(r, size=size, bold=header, color="FFFFFF" if header else INK)
    if header:
        shade(cell, DARK_BLUE)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    r = header.add_run("tissuePMHC 模型原理学习笔记")
    set_run_font(r, size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("从 E0 到 E29：共享结构、融合策略与 CNN 表示")
    set_run_font(r, size=8.5, color=MUTED)


def add_stage_table(doc):
    heading(doc, "一眼看懂：关键提升来自哪里", 1)
    paragraph(doc, "下面只列出真正改变研究路线、或者使性能发生明显提升的阶段。", size=10.5, color=MUTED, after=6)
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [700, 2250, 1250, 5160])
    headers = ["阶段", "模型", "Mean AUROC", "它解决了什么问题"]
    for cell, text in zip(table.rows[0].cells, headers):
        fill_table_cell(cell, text, header=True, center=(text != "它解决了什么问题"))
    set_repeat_table_header(table.rows[0])
    rows = [
        ("E0", "传统 one-hot 线性模型", "0.7558", "提供序列位置特征的基准，但无法自然学习局部 motif、任务共享或 HLA 层级结构。"),
        ("E2", "共享 peptide encoder + task heads", "0.7927", "让 44 个 tissue-HLA 任务共同学习 peptide 基础知识，同时保留每个任务自己的输出头。"),
        ("E8", "Global + HLA 双分支", "0.8050", "把全局规律与同一 HLA 内的特异规律分别建模，再进行软融合。"),
        ("E14", "Auxiliary global + plain HLA", "0.8116", "利用 tissue/HLA 辅助监督增强全局表示，同时避免干扰 HLA 分支的专门化。"),
        ("E15", "Task-rank fusion", "0.8130", "把两个分支变成 task 内排序再平均，避免概率尺度不一致。"),
        ("E17", "E14a 5-seed ensemble", "0.8263", "平均独立训练模型，削弱初始化与训练路径带来的偶然误差。"),
        ("E29", "Multi-kernel CNN E14a 3-seed", "0.8341", "显式提取 2、3、5 个氨基酸长度的局部 motif，并保留 E14 的强双分支结构。"),
    ]
    for stage, model, auroc, contribution in rows:
        cells = table.add_row().cells
        fill_table_cell(cells[0], stage, center=True)
        fill_table_cell(cells[1], model)
        fill_table_cell(cells[2], auroc, center=True)
        fill_table_cell(cells[3], contribution)
    paragraph(doc, "当前主结果：E29 3-seed ensemble，Mean AUROC 0.8341，Mean AUPRC 0.8228，Worst-10 Mean AUROC 0.7634。", size=10.5, bold=True, color=DARK_BLUE, before=5, after=5)


def build():
    doc = Document()
    configure_document(doc)

    p = paragraph(doc, "tissuePMHC 关键模型原理学习指南", size=24, bold=True, color=DARK_BLUE, before=10, after=3)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = paragraph(doc, "从传统线性模型到 E29 Multi-kernel CNN 3-seed ensemble", size=13, color=MUTED, after=16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    callout(
        doc,
        "先记住这一句话：",
        "本项目的主要提升不是来自把优化器变得更复杂，而是来自三件事：合理共享任务知识、保留 HLA 特异信息、让模型更好地识别 peptide 的局部序列模式。",
        fill=LIGHT_BLUE,
    )
    paragraph(doc, "阅读顺序建议：先看 E2、E8、E14，理解“结构”；再看 E15、E17，理解“融合与集成”；最后看 E29，理解当前最强模型为什么有效。", size=10.5, color=MUTED, after=10)

    add_stage_table(doc)

    heading(doc, "0. 任务到底是什么？", 1)
    paragraph(doc, "输入是一条长度为 9 的 peptide 序列。每个预测任务由“tissue + HLA allele”共同定义：模型要判断这条 peptide 在该 tissue-HLA 条件下是否更像正样本。项目共有 44 个任务。")
    callout(doc, "一个直观比喻", "把每个任务看成不同实验室的判定标准：它们都在看 peptide，但关注点不完全一样。有些规律跨所有实验室通用，有些规律只在某一类 HLA 内成立。", fill=LIGHT_GRAY)

    heading(doc, "1. E0：传统线性模型 - 只会看“每个位置是什么”", 1)
    paragraph(doc, "E0 将 9-mer 的每个位置和氨基酸展开成 one-hot 特征，再用 Logistic Regression 打分。它可以学习“第 2 位是某个氨基酸通常更好”这类独立位置效应。")
    labeled_paragraph(doc, "它擅长什么：", "快速、可解释，并能建立一个明确的传统基线。")
    labeled_paragraph(doc, "它缺少什么：", "它不擅长理解相邻氨基酸组合，也无法让不同任务共享经验。因此它更像一张位置-氨基酸打分表。")

    heading(doc, "2. E2：共享 peptide encoder - 44 个任务一起学习", 1)
    paragraph(doc, "E2 不再让 44 个任务各自训练完整模型，而是先让所有任务共享一个 peptide encoder。这个 encoder 将 peptide 压缩成一段“表示向量”，再交给各任务自己的输出头判断。")
    callout(doc, "结构示意", "peptide → 共享 encoder → 44 个 task-specific heads。共享 encoder 学通用 peptide 规律；每个 head 学本任务如何使用这些规律。", fill=LIGHT_BLUE)
    labeled_paragraph(doc, "为什么有效：", "单个任务的数据有限，但不同任务中的 peptide 模式有共性。共享 encoder 相当于让所有任务共同积累一套“读 peptide 的基础知识”。")
    labeled_paragraph(doc, "局限：", "E2 默认所有任务共享同一套核心表示，没有明确区分跨 HLA 的通用规律和某个 HLA 的专属规律。")

    heading(doc, "3. E8：Global + HLA 双分支 - 两个专家共同判断", 1)
    paragraph(doc, "E8 把共享策略拆成两条支路。Global branch 使用所有任务的数据，学习跨 tissue、跨 HLA 的通用模式；HLA branch 则按 HLA 分组，让同一 allele 下不同 tissue 的任务共享一个专门 encoder。")
    callout(doc, "结构示意", "同一条 peptide 同时进入 Global expert 和 HLA expert；前者问“整体规律是什么”，后者问“这个 HLA 特别偏好什么”；最后融合两者判断。", fill=LIGHT_BLUE)
    labeled_paragraph(doc, "为什么有效：", "这避免了二选一。若只用 global 模型，HLA motif 容易被稀释；若只按 HLA 分组，跨 HLA 的样本量和共性又被浪费。双分支同时保留两类信息。")
    labeled_paragraph(doc, "关键认识：", "这里的提升说明“如何共享”比单纯增加网络宽度更重要。")

    heading(doc, "4. E13 与 E14：辅助监督 - 让全局分支学得更有方向", 1)
    paragraph(doc, "E13 在主二分类任务之外，额外要求共享表示去预测 tissue 和 HLA。最终目标仍是正负分类；辅助任务的作用是迫使 encoder 形成更有组织的 peptide 表示。")
    labeled_paragraph(doc, "直观理解：", "普通训练只告诉模型“这题答对还是答错”；辅助训练还要求它说出“这条 peptide 更像与哪个 HLA、哪个 tissue 有关”。这会给 encoder 更多学习线索。")
    paragraph(doc, "E14 将这条思路与 E8 组合：Global branch 使用 tissue/HLA auxiliary supervision，HLA branch 保持普通训练。这个组合比给两条分支都加入辅助任务更好。")
    callout(doc, "为什么只增强 Global branch？", "Global branch 面对的任务范围最广，最需要借助 tissue/HLA 信息整理共享表示；HLA branch 已经按 allele 专门化，额外辅助约束可能反而限制其自由度。", fill=LIGHT_GRAY)

    heading(doc, "5. E15：Task-rank fusion - 先比较名次，再合并意见", 1)
    paragraph(doc, "两个分支都能输出概率，但概率数值不一定处于同一尺度。一个分支可能经常给出 0.80，另一个即使排得很准也只给出 0.55。直接平均概率，会让数值更极端的分支占更多话语权。")
    paragraph(doc, "E15 的做法是：在每个 task 内，先把每个分支的预测转换成百分位名次，再平均名次。它关注的是“这个样本在该 task 中排第几”，而不是“模型说它是 0.70 还是 0.90”。")
    labeled_paragraph(doc, "为什么适合 AUROC：", "AUROC 评价的核心就是正样本是否被排在负样本前面。因此使用 task 内 rank，更贴近最终评价目标。")

    heading(doc, "6. E17：多 seed ensemble - 让多个独立模型投票", 1)
    paragraph(doc, "即使模型结构、数据和超参数完全一样，随机初始化、batch 顺序和 dropout 也会让不同训练过程得到略有不同的模型。一个模型可能偶然错分某些样本，另一个模型未必会犯相同的错误。")
    callout(doc, "E17 的流程", "独立训练多个 E14a → Global 分支在 seed 间平均 → HLA 分支在 seed 间平均 → 最后进行 task-rank fusion。", fill=LIGHT_BLUE)
    labeled_paragraph(doc, "为什么有效：", "平均独立训练模型会保留多个模型共同认可的信号，同时削弱单个训练过程造成的偶然偏差。E17 5-seed 从 E14 的 0.8116 提升到 0.8263。")
    labeled_paragraph(doc, "为什么 checkpoint/SWA 没有同样成功：", "同一训练轨迹上的不同 checkpoint 太相似，错误也更相似；真正独立训练的 seed 才提供了更明显的多样性。")

    heading(doc, "7. E29：Multi-kernel CNN - 当前最强模型为什么更强", 1)
    paragraph(doc, "E14 的 encoder 先把 9 个位置的 amino-acid embedding 直接拼接，再交给 MLP。它可以学习位置效应，但没有显式结构来识别连续的局部氨基酸片段。")
    paragraph(doc, "E29 用三组一维卷积同时扫描 peptide：长度为 2 的卷积看二肽组合，长度为 3 的卷积看短 motif，长度为 5 的卷积看更长的局部片段。三种尺度的特征再拼接，交给后续网络。")
    callout(doc, "以 9-mer 为例", "若 peptide 是 SLYNTVATL，长度为 3 的卷积会重点观察 SLY、LYN、YNT、NTV 等连续片段。模型因而能直接学习“某种局部组合是否重要”，而不是只看每个位置的氨基酸。", fill=LIGHT_BLUE)
    labeled_paragraph(doc, "E29 的关键细节：", "卷积后的位置信息被保留，而不是只取一个全局最大值。对 HLA-I 9-mer，motif 出现在第 2 位附近还是第 9 位附近可能意义不同；保留位置能避免丢失这种锚定位置信息。")
    labeled_paragraph(doc, "为什么 E29 没有推倒重来：", "E29 只替换 peptide encoder，继续保留 E14 中已经证明有效的 Global auxiliary branch、HLA plain branch、task-rank fusion 和独立 seed ensemble。它是在强基线之上改善表示，而不是一次性改变所有因素。")
    paragraph(doc, "最终，E29 3-seed ensemble 的 Mean AUROC 为 0.8341，Mean AUPRC 为 0.8228，Worst-10 Mean AUROC 为 0.7634，超过此前 E17 5-seed 的结果。")

    heading(doc, "8. 最后把整条逻辑串起来", 1)
    callout(doc, "项目的核心结论", "先用 E2 让任务共享知识；再用 E8/E14 把全局规律、HLA 特异规律和辅助监督组合起来；用 E15 解决分支尺度问题；用 E17 降低训练随机性；最后由 E29 改善 peptide 的局部 motif 表示。", fill=LIGHT_BLUE)
    paragraph(doc, "因此，最有效的改进不是“越复杂越好”，而是每一步都针对一个明确瓶颈：共享范围、表示质量、分数可比性或训练方差。E26/E27 的负结果也很重要：如果候选模型本质上相似，再复杂的二层融合也无法凭空创造新信息。")
    heading(doc, "建议你记住的五个关键词", 2)
    paragraph(doc, "共享（E2）  ·  分层共享（E8）  ·  辅助监督（E14）  ·  排序融合与独立集成（E15/E17）  ·  局部 motif CNN（E29）", size=11, bold=True, color=DARK_BLUE, after=8)
    paragraph(doc, "注：本文用于帮助理解项目内部模型演进。所有性能数字均来自当前 closed-set standard split；它们不自动代表 peptide-disjoint、protein-disjoint、unseen-HLA 或外部数据上的泛化能力。", size=9.5, color=MUTED, italic=True, before=8, after=0)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
