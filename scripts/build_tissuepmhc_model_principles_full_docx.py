from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "results" / "tissuePMHC_model_principles_guide.docx"
FORMULA_DIR = ROOT / "results" / "tissuePMHC_model_principles_formulas"
BLACK = "000000"
DARK_GRAY = "404040"
MID_GRAY = "666666"
LIGHT_GRAY = "F2F2F2"
TABLE_GRAY = "D9E1F2"
PAGE_DXA = 9360


def font(run, size=11, bold=False, color=BLACK, italic=False, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), color)


def width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:tcW"))
    if node is None:
        node = OxmlElement("w:tcW")
        tc_pr.append(node)
    node.set(qn("w:w"), str(dxa))
    node.set(qn("w:type"), "dxa")


def geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    for tag, value in (("w:tblW", str(sum(widths))), ("w:tblInd", "120")):
        node = tbl_pr.first_child_found_in(tag)
        if node is None:
            node = OxmlElement(tag)
            tbl_pr.append(node)
        node.set(qn("w:w"), value)
        node.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for dxa in widths:
        child = OxmlElement("w:gridCol")
        child.set(qn("w:w"), str(dxa))
        grid.append(child)
    for row in table.rows:
        for cell, dxa in zip(row.cells, widths):
            width(cell, dxa)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def p(doc, text="", size=11, bold=False, color=BLACK, before=0, after=6, line=1.25, align=None, italic=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = line
    if align is not None:
        para.alignment = align
    if text:
        r = para.add_run(text)
        font(r, size, bold, color, italic)
    return para


def add_heading(doc, text, level=1):
    return p(doc, text, size={1: 16, 2: 13, 3: 12}[level], bold=True, color=BLACK,
             before={1: 18, 2: 14, 3: 10}[level], after={1: 9, 2: 7, 3: 5}[level])


def label_para(doc, label, body):
    para = p(doc, "", after=5)
    r = para.add_run(label)
    font(r, 11, True, DARK_GRAY)
    r = para.add_run(body)
    font(r, 11, False, BLACK)


def formula(doc, name, caption):
    image = FORMULA_DIR / f"{name}.png"
    if not image.is_file():
        raise FileNotFoundError(image)
    para = p(doc, "", before=2, after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    para.add_run().add_picture(str(image), width=Inches(4.8))
    p(doc, caption, size=9.5, color=MID_GRAY, after=6, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)


def callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    geometry(table, [PAGE_DXA])
    cell = table.cell(0, 0)
    shade(cell, LIGHT_GRAY)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(2)
    r = para.add_run(title)
    font(r, 11, True, DARK_GRAY)
    para = cell.add_paragraph()
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.line_spacing = 1.2
    r = para.add_run(body)
    font(r, 10.5, False, BLACK)
    p(doc, "", after=3)


def cell_text(cell, text, header=False, center=False):
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.08
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    r = para.add_run(text)
    font(r, 9.3, header, BLACK)
    if header:
        shade(cell, TABLE_GRAY)


def stage_table(doc):
    table = doc.add_table(rows=1, cols=4)
    geometry(table, [700, 2450, 1250, 4960])
    for cell, text in zip(table.rows[0].cells, ["阶段", "模型", "Mean AUROC", "核心贡献"]):
        cell_text(cell, text, header=True, center=text != "核心贡献")
    entries = [
        ("E0", "传统 one-hot Logistic Regression", "约 0.7558", "建立位置-氨基酸线性打分基线。"),
        ("E2", "共享 peptide encoder + task heads", "约 0.7927", "让 44 个任务共同学习 peptide 的通用表示。"),
        ("E8", "Global + HLA 双分支", "约 0.8050", "同时学习全局规律与 HLA 特异规律。"),
        ("E14", "Auxiliary global + plain HLA", "约 0.8116", "用辅助监督增强共享表示，并保留 HLA 专门化。"),
        ("E15", "Task-rank fusion", "约 0.8130", "消除两个分支概率尺度不一致的问题。"),
        ("E17", "E14a 5-seed ensemble", "约 0.8263", "用独立随机种子平均降低训练方差。"),
        ("E29", "Multi-kernel CNN E14a 3-seed", "0.8341", "显式提取局部 motif，并继续利用强双分支与 seed ensemble。"),
    ]
    for row in entries:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell_text(cells[i], text, center=i in (0, 2))
    p(doc, "", after=4)


def setup(doc):
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(1)
    sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(header.add_run("tissuePMHC 模型原理说明"), 8.5, False, MID_GRAY)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(footer.add_run("从 E0 到 E29：模型结构、融合与集成"), 8.5, False, MID_GRAY)


def build():
    doc = Document()
    setup(doc)

    p(doc, "tissuePMHC 关键模型原理说明", size=24, bold=True, color=BLACK, before=12, after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "从传统线性模型到 E29 Multi-kernel CNN 3-seed ensemble", size=13, color=DARK_GRAY, after=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    callout(doc, "阅读重点", "这份文档不试图罗列全部实验，而是解释那些真正改变性能主线的模型：E0、E2、E8、E13/E14、E15、E17 和 E29。请重点理解每一步解决了哪个旧模型的瓶颈。")

    add_heading(doc, "整体主线：模型是怎样一步一步变强的", 1)
    p(doc, "可以把整个研究过程理解成：模型先学会“多个任务共享知识”，再学会“保留 HLA 差异”，然后通过辅助监督、排序融合、随机种子集成和 CNN 局部模式提取逐步增强。")
    p(doc, "发展逻辑可以概括为：E0 传统线性模型 → E2 多任务共享编码器 → E8 双分支模型 → E14 辅助监督增强双分支 → E15 task-rank fusion → E17 多 seed 集成 → E29 Multi-kernel CNN 3-seed。", size=10.5, bold=True, color=DARK_GRAY, after=8)
    stage_table(doc)

    add_heading(doc, "先理解任务：模型究竟在预测什么？", 1)
    p(doc, "输入是一条长度固定为 9 的 peptide 序列。每个预测任务由 tissue 与 HLA allele 共同定义：模型需要判断这条 peptide 在给定 tissue-HLA 条件下是否更像正样本。项目的 standard split 包含 44 个任务。")
    callout(doc, "直观比喻", "把每个任务看成不同实验室的判定标准：它们都在看 peptide，但关注点不完全相同。有些规律跨所有实验室通用，有些规律只在某一类 HLA 内成立。")

    add_heading(doc, "1. E0：传统线性模型 - 只会看“每个位置是什么”", 1)
    p(doc, "E0 是最基础的序列打分器。每条 peptide 长度为 9，例如 SLYNTVATL。模型将“第几个位置是什么氨基酸”展开成 one-hot 特征，再使用 Logistic Regression 给出分数。")
    formula(doc, "e0_linear", "E0 的线性位置-氨基酸打分形式。")
    p(doc, "其中，位置 i 出现氨基酸 a 时，对预测的影响由对应权重决定。它能够学习“第 2 位是某个氨基酸通常更有利”这类独立位置效应。")
    label_para(doc, "它擅长什么：", "快速、可解释，并能建立一个明确的传统基线。")
    label_para(doc, "它缺少什么：", "它不擅长理解相邻氨基酸组合，也无法让不同任务共享经验。因此它更像一张位置-氨基酸打分表。")

    add_heading(doc, "2. E2：共享 peptide encoder - 44 个任务一起学习", 1)
    p(doc, "E2 不再让 44 个任务各自训练完整模型，而是先让所有任务共享一个 peptide encoder。这个 encoder 将 peptide 压缩成一段表示向量，再交给各任务自己的输出头判断。")
    formula(doc, "e2_shared", "共享表示与 task-specific head 的关系。")
    callout(doc, "结构示意", "peptide → 共享 encoder → 44 个 task-specific heads。共享 encoder 学通用 peptide 规律；每个 head 学本任务如何使用这些规律。")
    label_para(doc, "为什么有效：", "单个任务的数据有限，但不同任务中的 peptide 模式有共性。共享 encoder 相当于让所有任务共同积累一套“读 peptide 的基础知识”。")
    label_para(doc, "局限：", "E2 默认所有任务共享同一套核心表示，没有明确区分跨 HLA 的通用规律和某个 HLA 的专属规律。")

    add_heading(doc, "3. E8：Global + HLA 双分支 - 两个专家共同判断", 1)
    p(doc, "E8 把共享策略拆成两条支路。Global branch 使用所有任务的数据，学习跨 tissue、跨 HLA 的通用模式；HLA branch 则按 HLA 分组，让同一 allele 下不同 tissue 的任务共享一个专门 encoder。")
    callout(doc, "结构示意", "同一条 peptide 同时进入 Global expert 和 HLA expert。前者问“整体规律是什么”，后者问“这个 HLA 特别偏好什么”；最后融合两者判断。")
    formula(doc, "e8_fusion", "E8 的基础双分支平均融合。")
    label_para(doc, "为什么有效：", "这避免了二选一。若只用 global 模型，HLA motif 容易被稀释；若只按 HLA 分组，跨 HLA 的样本量和共性又被浪费。双分支同时保留两类信息。")
    label_para(doc, "关键认识：", "这里的提升说明“如何共享”比单纯增加网络宽度更重要。")

    add_heading(doc, "4. E13 与 E14：辅助监督 - 让全局分支学得更有方向", 1)
    p(doc, "E13 在主二分类任务之外，额外要求共享表示去预测 tissue 和 HLA。最终目标仍是正负分类；辅助任务的作用是迫使 encoder 形成更有组织的 peptide 表示。")
    label_para(doc, "直观理解：", "普通训练只告诉模型“这题答对还是答错”；辅助训练还要求它说出“这条 peptide 更像与哪个 HLA、哪个 tissue 有关”。这会给 encoder 更多学习线索。")
    p(doc, "E14 将这条思路与 E8 组合：Global branch 使用 tissue/HLA auxiliary supervision，HLA branch 保持普通训练。这个组合比给两条分支都加入辅助任务更好。")
    callout(doc, "为什么只增强 Global branch？", "Global branch 面对的任务范围最广，最需要借助 tissue/HLA 信息整理共享表示；HLA branch 已经按 allele 专门化，额外辅助约束可能反而限制其自由度。")

    add_heading(doc, "5. E15：Task-rank fusion - 先比较名次，再合并意见", 1)
    p(doc, "两个分支都能输出概率，但概率数值不一定处于同一尺度。一个分支可能经常给出 0.80，另一个即使排得很准也只给出 0.55。直接平均概率时，数值更极端的分支会产生更大影响。")
    p(doc, "E15 的做法是：在每个 task 内，先把每个分支的预测转换成百分位名次，再平均名次。它关注的是“这个样本在该 task 中排第几”，而不是“模型说它是 0.70 还是 0.90”。")
    formula(doc, "e15_rank", "E15 的 task-rank fusion：先转为 task 内排名，再平均。")
    label_para(doc, "为什么适合 AUROC：", "AUROC 评价的核心就是正样本是否被排在负样本前面。因此使用 task 内 rank，更贴近最终评价目标。")

    add_heading(doc, "6. E17：多 seed ensemble - 让多个独立模型投票", 1)
    p(doc, "即使模型结构、数据和超参数完全一样，随机初始化、batch 顺序和 dropout 也会让不同训练过程得到略有不同的模型。一个模型可能偶然错分某些样本，另一个模型未必会犯相同的错误。")
    formula(doc, "e17_seed", "多个独立 seed 的预测平均。")
    callout(doc, "E17 的流程", "独立训练多个 E14a → Global 分支在 seed 间平均 → HLA 分支在 seed 间平均 → 最后进行 task-rank fusion。")
    label_para(doc, "为什么有效：", "平均独立训练模型会保留多个模型共同认可的信号，同时削弱单个训练过程造成的偶然偏差。E17 5-seed 从 E14 的 0.8116 提升到 0.8263。")
    label_para(doc, "为什么 checkpoint/SWA 没有同样成功：", "同一训练轨迹上的不同 checkpoint 太相似，错误也更相似；真正独立训练的 seed 才提供了更明显的多样性。")

    add_heading(doc, "7. E29：Multi-kernel CNN - 当前最强模型为什么更强", 1)
    p(doc, "E14 的 peptide encoder 先把 9 个位置的 amino-acid embedding 直接拼接，再交给 MLP。它可以学习位置效应，但没有显式结构来识别连续的局部氨基酸片段。")
    p(doc, "E29 用三组一维卷积同时扫描 peptide：长度为 2 的卷积看二肽组合，长度为 3 的卷积看短 motif，长度为 5 的卷积看更长的局部片段。三种尺度的特征再拼接，交给后续网络。")
    callout(doc, "以 9-mer 为例", "若 peptide 是 SLYNTVATL，长度为 3 的卷积会重点观察 SLY、LYN、YNT、NTV 等连续片段。模型因而能直接学习“某种局部组合是否重要”，而不是只看每个位置的氨基酸。")
    label_para(doc, "E29 的关键细节：", "卷积后的位置信息被保留，而不是只取一个全局最大值。对 HLA-I 9-mer，motif 出现在第 2 位附近还是第 9 位附近可能意义不同；保留位置能避免丢失这种锚定位置信息。")
    label_para(doc, "为什么 E29 没有推倒重来：", "E29 只替换 peptide encoder，继续保留 E14 中已经证明有效的 Global auxiliary branch、HLA plain branch、task-rank fusion 和独立 seed ensemble。它是在强基线之上改善表示，而不是一次性改变所有因素。")
    formula(doc, "e29_gain", "E29 的核心：更好的局部 motif 表示与独立 seed 平均可以叠加。")
    p(doc, "最终，E29 3-seed ensemble 的 Mean AUROC 为 0.8341，Mean AUPRC 为 0.8228，Worst-10 Mean AUROC 为 0.7634，超过此前 E17 5-seed 的结果。")

    add_heading(doc, "最核心的研究结论", 1)
    p(doc, "这些实验最终指向了一个比较清晰的规律：有效提升主要来自合理的任务共享结构、互补的 global/HLA 分支、对 shared representation 有帮助的轻量辅助监督、与 AUROC 匹配的 task-rank fusion、真正独立的随机种子集成，以及符合 9-mer 局部 motif 特征的 CNN encoder。")
    p(doc, "相反，没有带来明显提升的方向包括：复杂的动态 loss weighting、仅在同一训练轨迹上平均 checkpoint、SWA、MC Dropout、同质候选上的复杂 stacking，以及过于复杂的 expert/gate 结构。")
    callout(doc, "一句话概括", "E2 让不同任务开始共享知识，E8 让共享变得有层次，E14 让共享表示获得辅助引导，E15 解决分支尺度差异，E17 降低训练随机性，E29 则真正增强了 peptide 局部 motif 的表示能力。")
    p(doc, "说明：本文用于理解项目内部模型演进。所有性能数字均来自当前 closed-set standard split；它们不自动代表 peptide-disjoint、protein-disjoint、unseen-HLA 或外部数据上的泛化能力。", size=9.5, color=MID_GRAY, italic=True, before=8, after=0)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
